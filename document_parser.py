"""
多格式文档解析器
支持 Word/Excel/PDF/MD/TXT 等格式
"""
import io
import logging
from pathlib import Path
from typing import Dict, Any, Optional, Tuple
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PyPDF2 import PdfReader
import markdown
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class DocumentParser:
    """多格式文档解析器"""

    SUPPORTED_FORMATS = {
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.pdf': 'application/pdf',
    }

    @staticmethod
    def get_file_extension(filename: str) -> str:
        """获取文件扩展名"""
        return Path(filename).suffix.lower()

    @staticmethod
    def is_supported(filename: str) -> bool:
        """检查文件格式是否支持"""
        ext = DocumentParser.get_file_extension(filename)
        return ext in DocumentParser.SUPPORTED_FORMATS

    @staticmethod
    def parse_txt(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """解析TXT文件"""
        try:
            # 尝试UTF-8解码
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            # 尝试GBK解码
            try:
                text = content.decode('gbk')
            except UnicodeDecodeError:
                # 尝试Latin1解码
                text = content.decode('latin1')

        metadata = {
            'filename': filename,
            'format': 'txt',
            'size': len(content),
            'lines': text.count('\n') + 1,
        }

        return text, metadata

    @staticmethod
    def parse_markdown(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """解析Markdown文件"""
        # 先解码为文本
        text = content.decode('utf-8')

        # 转换为HTML
        html = markdown.markdown(text)

        # 使用BeautifulSoup提取纯文本
        soup = BeautifulSoup(html, 'html.parser')
        plain_text = soup.get_text(separator='\n', strip=True)

        metadata = {
            'filename': filename,
            'format': 'markdown',
            'size': len(content),
            'lines': text.count('\n') + 1,
            'original_markdown': text[:500],  # 保留前500字符的原始Markdown
        }

        return plain_text, metadata

    @staticmethod
    def parse_docx(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """解析Word文档"""
        # 加载文档
        doc = DocxDocument(io.BytesIO(content))

        # 提取所有段落
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = '\n\n'.join(paragraphs)

        # 提取表格内容
        tables_text = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(' | '.join(row_data))
            if table_data:
                tables_text.append('\n'.join(table_data))

        if tables_text:
            text += '\n\n【表格内容】\n' + '\n\n'.join(tables_text)

        metadata = {
            'filename': filename,
            'format': 'docx',
            'size': len(content),
            'paragraphs': len(paragraphs),
            'tables': len(doc.tables),
        }

        return text, metadata

    @staticmethod
    def parse_xlsx(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """解析Excel文档"""
        # 加载工作簿
        wb = load_workbook(io.BytesIO(content), data_only=True)

        all_sheets_text = []

        # 遍历所有工作表
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]

            # 提取表格内容
            sheet_data = []
            for row in sheet.iter_rows(values_only=True):
                # 过滤空行
                if any(cell is not None for cell in row):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    sheet_data.append(row_text)

            if sheet_data:
                sheet_text = f'【工作表: {sheet_name}】\n' + '\n'.join(sheet_data)
                all_sheets_text.append(sheet_text)

        text = '\n\n'.join(all_sheets_text)

        metadata = {
            'filename': filename,
            'format': 'xlsx',
            'size': len(content),
            'sheets': len(wb.sheetnames),
            'sheet_names': wb.sheetnames,
        }

        return text, metadata

    @staticmethod
    def parse_pdf(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """解析PDF文档"""
        # 加载PDF
        pdf = PdfReader(io.BytesIO(content))

        # 提取所有页面的文本
        pages_text = []
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text.strip():
                pages_text.append(f'【第{i+1}页】\n{page_text}')

        text = '\n\n'.join(pages_text)

        metadata = {
            'filename': filename,
            'format': 'pdf',
            'size': len(content),
            'pages': len(pdf.pages),
        }

        return text, metadata

    @staticmethod
    def parse(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        解析文档（自动识别格式）

        Args:
            content: 文件内容（字节）
            filename: 文件名

        Returns:
            (文本内容, 元数据)
        """
        ext = DocumentParser.get_file_extension(filename)

        if not DocumentParser.is_supported(filename):
            raise ValueError(f"不支持的文件格式: {ext}")

        logger.info(f"Parsing {filename} (format: {ext}, size: {len(content)} bytes)")

        # 根据扩展名选择解析器
        if ext == '.txt':
            return DocumentParser.parse_txt(content, filename)
        elif ext == '.md':
            return DocumentParser.parse_markdown(content, filename)
        elif ext == '.docx':
            return DocumentParser.parse_docx(content, filename)
        elif ext == '.xlsx':
            return DocumentParser.parse_xlsx(content, filename)
        elif ext == '.pdf':
            return DocumentParser.parse_pdf(content, filename)
        else:
            raise ValueError(f"未实现的解析器: {ext}")


# 便捷函数
def parse_document(content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
    """便捷函数：解析文档"""
    return DocumentParser.parse(content, filename)


def is_format_supported(filename: str) -> bool:
    """便捷函数：检查格式是否支持"""
    return DocumentParser.is_supported(filename)


def get_supported_formats() -> Dict[str, str]:
    """便捷函数：获取支持的格式列表"""
    return DocumentParser.SUPPORTED_FORMATS.copy()


if __name__ == "__main__":
    # 测试代码
    print("支持的文档格式:")
    for ext, mime_type in get_supported_formats().items():
        print(f"  {ext}: {mime_type}")
