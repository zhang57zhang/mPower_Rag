#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档管理模块
支持多种文档格式的上传、删除、列表管理
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import hashlib
from datetime import datetime
import json

# 文档解析库
try:
    import docx  # python-docx for Word
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False

try:
    import openpyxl  # for Excel
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

try:
    import PyPDF2  # for PDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

logger = logging.getLogger(__name__)

class DocumentManager:
    """文档管理器"""
    
    def __init__(self, knowledge_base_dir: str = "knowledge_base"):
        self.knowledge_base_dir = Path(knowledge_base_dir)
        self.knowledge_base_dir.mkdir(exist_ok=True)
        
        # 文档元数据存储
        self.metadata_file = self.knowledge_base_dir / "documents_metadata.json"
        self.metadata = self._load_metadata()
        
        # 支持的文档格式
        self.supported_formats = {
            '.txt': self._parse_txt,
            '.md': self._parse_markdown,
            '.docx': self._parse_word if WORD_SUPPORT else None,
            '.xlsx': self._parse_excel if EXCEL_SUPPORT else None,
            '.pdf': self._parse_pdf if PDF_SUPPORT else None,
        }
    
    def _load_metadata(self) -> Dict:
        """加载文档元数据"""
        if self.metadata_file.exists():
            try:
                with open(self.metadata_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.error(f"Failed to load metadata: {e}")
                return {"documents": {}}
        return {"documents": {}}
    
    def _save_metadata(self):
        """保存文档元数据"""
        try:
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(self.metadata, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"Failed to save metadata: {e}")
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """计算文件哈希值"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hasher.update(chunk)
        return hasher.hexdigest()
    
    def _parse_txt(self, file_path: Path) -> str:
        """解析TXT文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _parse_markdown(self, file_path: Path) -> str:
        """解析Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _parse_word(self, file_path: Path) -> str:
        """解析Word文档"""
        if not WORD_SUPPORT:
            raise ImportError("python-docx not installed")
        
        doc = docx.Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return '\n'.join(text)
    
    def _parse_excel(self, file_path: Path) -> str:
        """解析Excel文档"""
        if not EXCEL_SUPPORT:
            raise ImportError("openpyxl not installed")
        
        wb = openpyxl.load_workbook(file_path)
        text = []
        
        for sheet in wb:
            text.append(f"## Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join([str(cell) if cell else '' for cell in row])
                text.append(row_text)
        
        return '\n'.join(text)
    
    def _parse_pdf(self, file_path: Path) -> str:
        """解析PDF文档"""
        if not PDF_SUPPORT:
            raise ImportError("PyPDF2 not installed")
        
        text = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text())
        
        return '\n'.join(text)
    
    def upload_document(
        self, 
        file_path: Path, 
        category: str = "general",
        metadata: Optional[Dict] = None
    ) -> Dict[str, Any]:
        """
        上传文档
        
        Args:
            file_path: 文件路径
            category: 文档分类
            metadata: 额外元数据
        
        Returns:
            上传结果
        """
        # 检查文件格式
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_formats:
            return {
                "success": False,
                "error": f"Unsupported format: {file_ext}. Supported: {list(self.supported_formats.keys())}"
            }
        
        parser = self.supported_formats[file_ext]
        if parser is None:
            return {
                "success": False,
                "error": f"Parser not available for {file_ext}. Please install required library."
            }
        
        try:
            # 解析文档
            content = parser(file_path)
            
            # 计算文件哈希
            file_hash = self._calculate_file_hash(file_path)
            
            # 生成文档ID
            doc_id = f"doc_{file_hash[:12]}"
            
            # 创建文档文件名
            doc_filename = f"{doc_id}{file_ext}"
            doc_path = self.knowledge_base_dir / doc_filename
            
            # 复制文件到知识库
            import shutil
            shutil.copy2(file_path, doc_path)
            
            # 保存解析后的文本
            text_path = self.knowledge_base_dir / f"{doc_id}.txt"
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # 更新元数据
            doc_metadata = {
                "id": doc_id,
                "filename": file_path.name,
                "category": category,
                "format": file_ext,
                "hash": file_hash,
                "size": file_path.stat().st_size,
                "upload_time": datetime.now().isoformat(),
                "content_length": len(content),
                "custom_metadata": metadata or {}
            }
            
            self.metadata["documents"][doc_id] = doc_metadata
            self._save_metadata()
            
            logger.info(f"Document uploaded: {doc_id} - {file_path.name}")
            
            return {
                "success": True,
                "doc_id": doc_id,
                "filename": file_path.name,
                "category": category,
                "content_length": len(content),
                "content": content  # 返回内容用于向量化
            }
            
        except Exception as e:
            logger.error(f"Failed to upload document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def delete_document(self, doc_id: str) -> Dict[str, Any]:
        """
        删除文档
        
        Args:
            doc_id: 文档ID
        
        Returns:
            删除结果
        """
        if doc_id not in self.metadata["documents"]:
            return {
                "success": False,
                "error": f"Document not found: {doc_id}"
            }
        
        try:
            doc_metadata = self.metadata["documents"][doc_id]
            
            # 删除文件
            doc_ext = doc_metadata["format"]
            doc_file = self.knowledge_base_dir / f"{doc_id}{doc_ext}"
            text_file = self.knowledge_base_dir / f"{doc_id}.txt"
            
            if doc_file.exists():
                doc_file.unlink()
            if text_file.exists():
                text_file.unlink()
            
            # 删除元数据
            del self.metadata["documents"][doc_id]
            self._save_metadata()
            
            logger.info(f"Document deleted: {doc_id}")
            
            return {
                "success": True,
                "doc_id": doc_id,
                "filename": doc_metadata["filename"]
            }
            
        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def list_documents(
        self, 
        category: Optional[str] = None,
        format: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        列出文档
        
        Args:
            category: 按分类过滤
            format: 按格式过滤
        
        Returns:
            文档列表
        """
        documents = []
        
        for doc_id, doc_metadata in self.metadata["documents"].items():
            # 过滤条件
            if category and doc_metadata["category"] != category:
                continue
            if format and doc_metadata["format"] != format:
                continue
            
            documents.append({
                "id": doc_id,
                "filename": doc_metadata["filename"],
                "category": doc_metadata["category"],
                "format": doc_metadata["format"],
                "size": doc_metadata["size"],
                "upload_time": doc_metadata["upload_time"],
                "content_length": doc_metadata["content_length"]
            })
        
        return documents
    
    def get_document(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """
        获取文档详情
        
        Args:
            doc_id: 文档ID
        
        Returns:
            文档详情
        """
        if doc_id not in self.metadata["documents"]:
            return None
        
        doc_metadata = self.metadata["documents"][doc_id]
        
        # 读取文档内容
        text_file = self.knowledge_base_dir / f"{doc_id}.txt"
        content = None
        if text_file.exists():
            with open(text_file, 'r', encoding='utf-8') as f:
                content = f.read()
        
        return {
            "id": doc_id,
            "filename": doc_metadata["filename"],
            "category": doc_metadata["category"],
            "format": doc_metadata["format"],
            "size": doc_metadata["size"],
            "upload_time": doc_metadata["upload_time"],
            "content_length": doc_metadata["content_length"],
            "content": content,
            "metadata": doc_metadata["custom_metadata"]
        }
    
    def get_supported_formats(self) -> Dict[str, bool]:
        """获取支持的文档格式"""
        return {
            "txt": True,
            "md": True,
            "docx": WORD_SUPPORT,
            "xlsx": EXCEL_SUPPORT,
            "pdf": PDF_SUPPORT
        }
