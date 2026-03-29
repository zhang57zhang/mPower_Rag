"""
文档加载器
支持多格式文档加载、分块和管理
"""
import os
import hashlib
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, field

from langchain_core.documents import Document
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
)

logger = logging.getLogger(__name__)


@dataclass
class ChunkMetadata:
    """分块元数据"""
    source: str
    chunk_index: int
    total_chunks: int
    content_hash: str
    file_type: str
    file_size: int
    created_at: str = field(default_factory=lambda: "")
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "source": self.source,
            "chunk_index": self.chunk_index,
            "total_chunks": self.total_chunks,
            "content_hash": self.content_hash,
            "file_type": self.file_type,
            "file_size": self.file_size,
            "created_at": self.created_at,
        }


class DocumentManager:
    """
    文档管理器
    
    负责：
    1. 多格式文档加载（TXT, MD, DOCX, PDF, XLSX等）
    2. 智能分块（保留语义边界）
    3. 元数据提取
    4. 去重管理
    """
    
    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
        separators: Optional[List[str]] = None,
    ):
        """
        初始化文档管理器
        
        Args:
            chunk_size: 分块大小（字符数）
            chunk_overlap: 分块重叠大小
            separators: 分隔符优先级列表
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # 默认分隔符（按语义边界分割）
        self.separators = separators or [
            "\n\n",      # 段落
            "\n",        # 行
            "。",        # 中文句号
            "．",        # 中文句号（全角）
            ".",         # 英文句号
            "！",        # 中文感叹号
            "!",         # 英文感叹号
            "？",        # 中文问号
            "?",         # 英文问号
            "；",        # 中文分号
            ";",         # 英文分号
            "，",        # 中文逗号
            ",",         # 英文逗号
            " ",         # 空格
            "",          # 字符级
        ]
        
        # 初始化文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=self.separators,
            length_function=len,
        )
        
        # Markdown 分割器
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "header_1"),
                ("##", "header_2"),
                ("###", "header_3"),
            ]
        )
        
        # 内容哈希索引（用于去重）
        self._content_hashes: Dict[str, str] = {}
        
        logger.info(f"DocumentManager initialized with chunk_size={chunk_size}, overlap={chunk_overlap}")
    
    def load_and_split(
        self,
        file_path: str,
        encoding: str = "utf-8",
    ) -> List[Document]:
        """
        加载文件并分块
        
        Args:
            file_path: 文件路径
            encoding: 文件编码
            
        Returns:
            分块后的文档列表
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # 获取文件类型
        file_type = path.suffix.lower()
        
        # 根据文件类型选择加载器
        if file_type == ".md":
            return self._load_markdown(file_path, encoding)
        elif file_type in [".txt", ".text"]:
            return self._load_text(file_path, encoding)
        elif file_type == ".docx":
            return self._load_docx(file_path)
        elif file_type == ".pdf":
            return self._load_pdf(file_path)
        elif file_type in [".xlsx", ".xls"]:
            return self._load_excel(file_path)
        else:
            # 默认作为文本处理
            return self._load_text(file_path, encoding)
    
    def _load_text(
        self,
        file_path: str,
        encoding: str = "utf-8",
    ) -> List[Document]:
        """加载纯文本文件"""
        with open(file_path, "r", encoding=encoding) as f:
            content = f.read()
        
        return self._split_content(
            content=content,
            source=file_path,
            file_type=".txt",
        )
    
    def _load_markdown(
        self,
        file_path: str,
        encoding: str = "utf-8",
    ) -> List[Document]:
        """加载 Markdown 文件（保留标题结构）"""
        with open(file_path, "r", encoding=encoding) as f:
            content = f.read()
        
        # 先按标题分割
        try:
            md_chunks = self.markdown_splitter.split_text(content)
        except Exception:
            # 如果 Markdown 分割失败，使用普通分割
            return self._split_content(
                content=content,
                source=file_path,
                file_type=".md",
            )
        
        # 对每个 Markdown 块进行进一步分割
        documents = []
        chunk_index = 0
        
        for md_chunk in md_chunks:
            # 如果块太大，进一步分割
            if len(md_chunk.page_content) > self.chunk_size:
                sub_chunks = self.text_splitter.split_text(md_chunk.page_content)
                for sub_content in sub_chunks:
                    doc = Document(
                        page_content=sub_content,
                        metadata={
                            **md_chunk.metadata,
                            "source": file_path,
                            "file_type": ".md",
                            "chunk_index": chunk_index,
                        }
                    )
                    documents.append(doc)
                    chunk_index += 1
            else:
                doc = Document(
                    page_content=md_chunk.page_content,
                    metadata={
                        **md_chunk.metadata,
                        "source": file_path,
                        "file_type": ".md",
                        "chunk_index": chunk_index,
                    }
                )
                documents.append(doc)
                chunk_index += 1
        
        # 更新总块数
        total = len(documents)
        for doc in documents:
            doc.metadata["total_chunks"] = total
        
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("Please install python-docx: pip install python-docx")
        
        doc = DocxDocument(file_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        content = "\n\n".join(paragraphs)
        
        return self._split_content(
            content=content,
            source=file_path,
            file_type=".docx",
        )
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载 PDF 文档"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF not installed")
            raise ImportError("Please install PyMuPDF: pip install pymupdf")
        
        doc = fitz.open(file_path)
        
        # 提取每页文本
        pages = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            if text.strip():
                pages.append(f"[Page {page_num + 1}]\n{text}")
        
        content = "\n\n".join(pages)
        doc.close()
        
        return self._split_content(
            content=content,
            source=file_path,
            file_type=".pdf",
        )
    
    def _load_excel(self, file_path: str) -> List[Document]:
        """加载 Excel 文件"""
        try:
            import pandas as pd
        except ImportError:
            logger.error("pandas not installed")
            raise ImportError("Please install pandas: pip install pandas openpyxl")
        
        # 读取 Excel
        df = pd.read_excel(file_path, sheet_name=None)
        
        # 合并所有工作表
        contents = []
        for sheet_name, sheet_df in df.items():
            contents.append(f"## Sheet: {sheet_name}\n")
            contents.append(sheet_df.to_string(index=False))
            contents.append("\n")
        
        content = "\n".join(contents)
        
        return self._split_content(
            content=content,
            source=file_path,
            file_type=".xlsx",
        )
    
    def _split_content(
        self,
        content: str,
        source: str,
        file_type: str,
    ) -> List[Document]:
        """
        分割内容为块
        
        Args:
            content: 文本内容
            source: 来源文件
            file_type: 文件类型
            
        Returns:
            文档块列表
        """
        # 使用分割器分割
        chunks = self.text_splitter.split_text(content)
        
        # 创建文档对象
        documents = []
        total_chunks = len(chunks)
        file_size = len(content)
        
        for idx, chunk in enumerate(chunks):
            # 计算内容哈希（用于去重）
            content_hash = self._compute_hash(chunk)
            
            # 检查是否重复
            if content_hash in self._content_hashes:
                logger.debug(f"Skipping duplicate chunk {idx} from {source}")
                continue
            
            # 记录哈希
            self._content_hashes[content_hash] = source
            
            # 创建文档
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "file_type": file_type,
                    "chunk_index": idx,
                    "total_chunks": total_chunks,
                    "content_hash": content_hash,
                    "file_size": file_size,
                }
            )
            documents.append(doc)
        
        logger.info(f"Split {source} into {len(documents)} chunks (skipped {total_chunks - len(documents)} duplicates)")
        
        return documents
    
    def _compute_hash(self, content: str) -> str:
        """计算内容哈希"""
        # 标准化内容（去除空白）
        normalized = " ".join(content.split())
        return hashlib.md5(normalized.encode()).hexdigest()
    
    def check_duplicate(self, content: str) -> Optional[str]:
        """
        检查内容是否重复
        
        Args:
            content: 待检查内容
            
        Returns:
            如果重复，返回已存在的来源；否则返回 None
        """
        content_hash = self._compute_hash(content)
        return self._content_hashes.get(content_hash)
    
    def clear_hashes(self):
        """清空哈希索引"""
        self._content_hashes.clear()
        logger.info("Content hash index cleared")


# 全局文档管理器实例
_document_manager: Optional[DocumentManager] = None


def get_document_manager(
    chunk_size: int = 512,
    chunk_overlap: int = 50,
    **kwargs
) -> DocumentManager:
    """
    获取文档管理器（单例）
    
    Args:
        chunk_size: 分块大小
        chunk_overlap: 分块重叠
        **kwargs: 其他参数
        
    Returns:
        文档管理器实例
    """
    global _document_manager
    
    if _document_manager is None:
        _document_manager = DocumentManager(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            **kwargs
        )
    
    return _document_manager
